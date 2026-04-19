"""
Final Report Generator - AI Recommendations Summary
Generates the summary report with specific segments requested by the user.
"""
import logging
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

from ..config import TEMPLATE_DIR
from ..utils.rules_manager import RulesManager

logger = logging.getLogger(__name__)

class FinalReportGenerator:
    """Generates the summary Final Report (DOCX)"""
    
    def __init__(self, output_path: str, font_option: str = 'times', language: str = 'vi'):
        self.output_path = output_path
        self.font_option = font_option.lower()
        self.language = language.lower() # 'vi' or 'en'
        
        # Select template
        template_name = 'timenr_template_report.docx' if self.font_option == 'times' else 'calibri_template_report.docx'
        template_path = TEMPLATE_DIR / template_name
        
        try:
            if template_path.exists():
                self.doc = Document(str(template_path))
                logger.info(f"Initialized Final Report with template: {template_path}")
            else:
                logger.warning(f"Template {template_path} not found. Using default document.")
                self.doc = Document()
        except Exception as e:
            logger.error(f"Error loading report template: {e}")
            self.doc = Document()

        self.font_name = 'Times New Roman' if self.font_option == 'times' else 'Calibri'
        self._setup_styles()

    def _setup_styles(self):
        """Configure default fonts and styles"""
        try:
            style = self.doc.styles['Normal']
            style.font.name = self.font_name
            style.font.size = Pt(11)
            
            def force_font(style_obj, name):
                rPr = style_obj._element.rPr
                if rPr is None:
                    rPr = style_obj._element.make_element(qn('w:rPr'))
                    style_obj._element.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}" w:eastAsia="{name}"/>')
                    rPr.append(rFonts)
                else:
                    rFonts.set(qn('w:ascii'), name)
                    rFonts.set(qn('w:hAnsi'), name)
            
            force_font(style, self.font_name)
            
            headings = [('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11)]
            for name, size in headings:
                if name in self.doc.styles:
                    h = self.doc.styles[name]
                    h.font.name = self.font_name
                    h.font.size = Pt(size)
                    h.font.bold = True
                    h.font.color.rgb = RGBColor(0, 0, 0)
                    force_font(h, self.font_name)
        except Exception as e:
            logger.warning(f"Report style setup failed: {e}")

    def generate(self, parsed_data: Dict[str, Any], db_name: str = None) -> bool:
        """Main entry point to build the report"""
        try:
            # Use all nodes for aggregated evaluation
            nodes = parsed_data.get('nodes', [])
            if not nodes:
                logger.error("No node data found for report generation")
                return False
            
            node1 = nodes[0] # Primary node for general info
            
            if not db_name:
                db_name = str(parsed_data.get('db_name', 'UNKNOWN')).upper()
            else:
                db_name = db_name.upper()
            
            # --- Remove potential blank line at start ---
            if self.doc.paragraphs and not self.doc.paragraphs[0].text.strip():
                p = self.doc.paragraphs[0]._element
                p.getparent().remove(p)

            # --- 3.1 DB NAME ---
            self.doc.add_heading(db_name, level=2)
            
            # --- 3.1.1 General Information (Node 1 only) ---
            self._add_general_info_section(node1)
            
            # --- Load and Evaluate Rules (All Nodes Aggregated) ---
            rules = RulesManager.load_rules()
            findings = self._evaluate_aggregated_rules(nodes, rules)
            
            # --- 3.1.2 Evaluation / Summary ---
            self._add_evaluation_section(findings)
            
            # --- Page Break ---
            self.doc.add_page_break()
            
            # --- 3.1.3 Recommendation ---
            self._add_recommendation_section(findings)
            
            # --- 3.1.4 Appendix Reference ---
            self._add_appendix_ref_section(db_name)
            
            # Save
            self.doc.save(self.output_path)
            logger.info(f"Final report saved to: {self.output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to generate final report: {e}")
            return False

    def _add_general_info_section(self, node_data: Dict[str, Any]):
        title = "Thông tin tổng quan" if self.language == 'vi' else "General information"
        self.doc.add_heading(title, level=3)
        
        # Table with 3 columns, 12 rows (header + 11 data?) 
        # User said: "bảng với 3 cột và 12 dòng, header row là 'Mục', 'Thông tin', 'Thông tin (Báo cáo trước đây)'"
        # 1 header + 11 data = 12 rows.
        table = self.doc.add_table(rows=12, cols=3)
        table.style = 'Table Grid'
        table.allow_autofit = False
        
        # Header
        headers = ["Mục", "Thông tin", "Thông tin (Báo cáo trước đây)"] if self.language == 'vi' else \
                  ["ITEM", "VALUE", "VALUE (PREVIOUS REPORT)"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        self._format_header_row(table.rows[0])
        
        # Fixed widths: 1 - 4.4 cm, 2 - 6.0 cm, 3 - 6.0 cm (Total 16.4cm)
        self._set_column_widths(table, [Cm(4.4), Cm(6.0), Cm(6.0)])
        
        # Data from REPORT_DETAILS (Node 1)
        report_details = node_data.get('database_info', {}).get('REPORT_DETAILS', [])
        # report_details[0] is likely header, data starts from index 1
        data_rows = report_details[1:] if len(report_details) > 1 else []
        
        for i in range(1, 12):
            idx = i - 1
            if idx < len(data_rows):
                row_data = data_rows[idx]
                if len(row_data) >= 2:
                    table.rows[i].cells[0].text = str(row_data[0])
                    table.rows[i].cells[1].text = str(row_data[1])
            
            # Column 3 is always empty
            table.rows[i].cells[2].text = ""

            # Formatting: Vertical center for all, Column 1 is Left-aligned
            for j in range(3):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                # User specifically asked: "cột 1 ở mục 1.1.1. cần set align center left"
                # Center vertical is already handled globally if I add it to _set_column_widths or here
                cell = table.rows[i].cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = align

    def _evaluate_aggregated_rules(self, nodes: List[Dict[str, Any]], rules: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Evaluate database data across ALL nodes and aggregate findings"""
        findings = []
        
        def sort_key(s):
            parts = s.split('.')
            return [int(p) if p.isdigit() else p for p in parts]
        
        sorted_rule_ids = sorted(rules.keys(), key=sort_key)
        
        for rid in sorted_rule_ids:
            rule = rules[rid]
            category = rule.get('category')
            target = rule.get('target')
            col_name = rule.get('column', '').upper()
            condition = rule.get('condition')
            threshold = rule.get('threshold')
            
            all_node_matches = []
            unique_names = set() # To prevent duplicates across rows/nodes
            
            for node_idx, node_data in enumerate(nodes):
                # --- Logic by Category ---
                if category == 'DB_INFO':
                    db_info = node_data.get('database_info', {})
                    table = db_info.get(target, [])
                    if not table or len(table) < 2: continue
                    headers = table[0]
                    
                    if condition == 'oracle_version_patch_age':
                        # Special Rule: Standardized Oracle Version & Patch Age Check
                        # Find column index
                        content_idx = next((i for i, h in enumerate(headers) if col_name in h.upper()), None)
                        if content_idx is None: continue
                        
                        full_text = ""
                        for row in table[1:]:
                            if len(row) > content_idx:
                                full_text += str(row[content_idx]) + "\n"
                        
                        import re
                        from datetime import datetime
                        
                        # 1. Detect Version
                        version_match = re.search(r'Database Release Update\s*:\s*(\d+)\.', full_text)
                        if version_match:
                            major_version = int(version_match.group(1))
                            if major_version < 19:
                                finding_name = f"Oracle Version {major_version} (Legacy)"
                                if finding_name not in unique_names:
                                    all_node_matches.append({"name": finding_name, "value": "Legacy", "branch": "version_old"})
                                    unique_names.add(finding_name)
                            else:
                                # Check Patch Age for 19c+
                                ru_patch_id_match = re.search(r'Database Release Update\s*:\s*.*\((\d+)\)', full_text)
                                if ru_patch_id_match:
                                    ru_patch_id = ru_patch_id_match.group(1)
                                    date_pattern = rf'Patch\s+{ru_patch_id}.*?applied on.*?(?:[a-zA-Z]{{3,}}\s+)?([a-zA-Z]{{3}})\s+(\d{{1,2}}).*?(\d{{4}})'
                                    date_match = re.search(date_pattern, full_text, re.DOTALL | re.IGNORECASE)
                                    if date_match:
                                        mon_str, day_str, year_str = date_match.groups()
                                        mon_key = mon_str.capitalize()
                                        months_map = {'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6, 
                                                     'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12}
                                        
                                        try:
                                            patch_date = datetime(int(year_str), months_map.get(mon_key, 1), int(day_str))
                                            current_date = datetime(2026, 4, 19)
                                            diff_months = (current_date.year - patch_date.year) * 12 + (current_date.month - patch_date.month)
                                            
                                            if diff_months >= 24: branch = "age_2y"
                                            elif diff_months >= 12: branch = "age_1y"
                                            elif diff_months >= 6: branch = "age_6m"
                                            else: branch = None
                                            
                                            if branch:
                                                finding_name = f"Oracle 19c Patch Age ({diff_months} months)"
                                                if finding_name not in unique_names:
                                                    all_node_matches.append({"name": finding_name, "value": f"Applied {mon_key} {year_str}", "branch": branch})
                                                    unique_names.add(finding_name)
                                        except: pass
                        continue # End DB_INFO for this rule

                    # Fuzzy column matching: Remove spaces and symbols to find partial matches
                    def normalize_header(h): return "".join(c for c in str(h).upper() if c.isalnum())
                    norm_col_name = normalize_header(col_name)
                    
                    col_idx = None
                    for i, h in enumerate(headers):
                        if norm_col_name in normalize_header(h):
                            col_idx = i
                            break
                    
                    if col_idx is None: continue
                    
                    for row in table[1:]:
                        if len(row) <= col_idx: continue
                        # Clean numeric strings (remove %, GB, MB, commas)
                        raw_val = str(row[col_idx]).strip()
                        val_str = "".join(c for c in raw_val if c.isdigit() or c == '.')
                        
                        try:
                            val = float(val_str) if '.' in val_str else int(val_str)
                        except:
                            val = raw_val # Fallback to string
                            
                        is_violation = False
                        if condition == ">": 
                            is_violation = (isinstance(val, (int, float)) and isinstance(threshold, (int, float)) and val > threshold)
                        elif condition == "<": 
                            is_violation = (isinstance(val, (int, float)) and isinstance(threshold, (int, float)) and val < threshold)
                        elif condition == "==": 
                            is_violation = (str(val) == str(threshold))
                        elif condition == "!=": 
                            is_violation = (str(val) != str(threshold))
                        elif condition == "contains": 
                            is_violation = (str(threshold).lower() in str(val).lower())
                        elif condition == "count":
                            # For count, we check the total items in the table
                            is_violation = (len(table)-1 > (int(threshold) if threshold != "" else 0))
                            if is_violation and "Aggregated" not in unique_names:
                                all_node_matches.append({"name": f"{len(table)-1} items", "value": len(table)-1})
                                unique_names.add("Aggregated")
                            break

                        if is_violation:
                            item_name = row[0]
                            if item_name not in unique_names:
                                all_node_matches.append({"name": item_name, "value": val})
                                unique_names.add(item_name)
                
                elif category == 'LOG':
                    # Fix: The parser data structure uses 'alerts' key, and internal key is 'error_code'
                    alert_data = node_data.get('alerts', {})
                    errors_list = alert_data.get('alerts', [])
                    
                    import re
                    ora_pattern = re.compile(r'(ORA-\d+)')
                    
                    if condition == "contains_any":
                        for err in errors_list:
                            raw_code = str(err.get('error_code', ''))
                            # Failsafe: Clean the code again in the generator
                            match = ora_pattern.search(raw_code)
                            code = match.group(1) if match else raw_code
                            
                            # Check if code matches any of the threshold strings
                            t_list = threshold if isinstance(threshold, list) else [threshold]
                            if any(t.lower() in code.lower() for t in t_list):
                                if code not in unique_names:
                                    all_node_matches.append({"name": code, "value": "Found"})
                                    unique_names.add(code)


            # --- Create Aggregated Finding ---
            if all_node_matches:
                # Limit to top 10 unique items
                display_matches = all_node_matches[:10]
                item_names = ", ".join([str(m['name']) for m in display_matches])
                if len(all_node_matches) > 10: 
                    item_names += f" and {len(all_node_matches)-10} more"
                
                best_match = all_node_matches[0]
                sample_val = best_match['value']
                
                # Handle Standardized Branches for oracle_version_patch_age
                branch = best_match.get('branch')
                if condition == 'oracle_version_patch_age' and branch:
                    if branch == "version_old":
                        rec_vi = "Phiên bản hiện tại đã cũ. Xem xét kế hoạch nâng cấp CSDL lên phiên bản mới hơn như Oracle 19c để được hỗ trợ tốt hơn từ phía hãng Oracle, giảm thiểu tỉ lệ phát sinh lỗi, nâng cao hiệu suất và có thể sử dụng các tính năng mới."
                        rec_en = "Current database version is no longer supported from Oracle. It is suggested to upgrade to version 19c and apply latest patch for Database to minimize the impact of known bugs or security threats."
                        sev_vi, sev_en = "Nghiêm trọng", "Critical"
                        risk_vi = "CSDL không được đảm bảo tính bảo mật ở mức cao nhất. Ngoài ra, có thể gặp bugs ảnh hưởng tới hoạt động của CSDL."
                        risk_en = "System is running on an outdated version that is no longer in primary support."
                    elif branch == "age_2y":
                        rec_vi = "Bản vá của cơ sở dữ liệu đã cũ quá 2 năm. Xem xét nâng cấp lên bản patch mới nhất để giảm thiểu rủi ro bảo mật và tránh phát sinh lỗi."
                        rec_en = "Patch update is more than 2 years out-of-date. It is suggested to apply latest patch for Database to minimize the impact of known bugs or security threats."
                        sev_vi, sev_en = "Nghiêm trọng", "Critical"
                        risk_vi = "CSDL không được đảm bảo tính bảo mật ở mức cao nhất. Ngoài ra, có thể gặp bugs ảnh hưởng tới hoạt động của CSDL."
                        risk_en = "Database may be unstable."
                    elif branch == "age_1y":
                        rec_vi = "Bản vá của cơ sở dữ liệu đã cũ. Xem xét nâng cấp lên bản patch mới nhất để giảm thiểu rủi ro bảo mật và tránh phát sinh lỗi."
                        rec_en = "Patch update is more than 1 year out-of-date. It is suggested to apply latest patch for Database to minimize the impact of known bugs or security threats."
                        sev_vi, sev_en = "Cao", "High"
                        risk_vi = "CSDL không được đảm bảo tính bảo mật ở mức cao nhất. Ngoài ra, có thể gặp bugs ảnh hưởng tới hoạt động của CSDL."
                        risk_en = "Database may be unstable."
                    elif branch == "age_6m":
                        rec_vi = "Bản vá của cơ sở dữ liệu đã cũ. Xem xét nâng cấp lên bản patch mới nhất để giảm thiểu rủi ro bảo mật và tránh phát sinh lỗi."
                        rec_en = "Patch update is almost 1 year out-of-date. It is suggested to apply latest patch for Database to minimize the impact of known bugs or security threats."
                        sev_vi, sev_en = "Thấp", "Low"
                        risk_vi = "CSDL không được đảm bảo tính bảo mật ở mức cao nhất. Ngoài ra, có thể gặp bugs ảnh hưởng tới hoạt động của CSDL."
                        risk_en = "Database may be unstable."
                    else:
                        sev_vi, sev_en = "Thấp", "Low"
                        rec_vi, rec_en, risk_vi, risk_en = "", "", "", ""

                    finding = {
                        "id": rid,
                        "title": rule.get('title'),
                        "item_name": item_names,
                        "value": sample_val,
                        "rec_vi": rec_vi, "rec_en": rec_en,
                        "risk_vi": risk_vi, "risk_en": risk_en,
                        "severity_vi": sev_vi, "severity_en": sev_en,
                        "appendix_ref": rule.get('appendix_ref', f"Phụ lục {rid}")
                    }
                else:
                    # Default handle for other rules
                    finding = {
                        "id": rid,
                        "title": rule.get('title'),
                        "item_name": item_names,
                        "value": sample_val,
                        "rec_vi": rule.get('rec_vi', '').replace("{item_name}", item_names).replace("{value}", str(sample_val)),
                        "rec_en": rule.get('rec_en', '').replace("{item_name}", item_names).replace("{value}", str(sample_val)),
                        "risk_vi": rule.get('risk_vi', '').replace("{item_name}", item_names).replace("{value}", str(sample_val)),
                        "risk_en": rule.get('risk_en', '').replace("{item_name}", item_names).replace("{value}", str(sample_val)),
                        "severity_vi": rule.get('severity_vi', 'Trung bình'),
                        "severity_en": rule.get('severity_en', 'Medium'),
                        "appendix_ref": rule.get('appendix_ref', f"Phụ lục {rid}")
                    }
                findings.append(finding)
                
        return findings

    def _add_evaluation_section(self, findings: List[Dict[str, Any]]):
        if self.language == 'vi':
            title = "Đánh giá"
            headers = ["STT", "Kiểm tra", "Đánh giá", "Đánh giá (Báo cáo trước đây)"]
            items = [
                "Kiểm tra trạng thái",
                "Kiểm tra hiệu năng - CPU",
                "Kiểm tra hiệu năng - Bộ nhớ",
                "Kiểm tra hiệu năng - Storage (còn trống) trên mỗi vùng đĩa",
                "Kiểm tra hiệu năng - Storage (I/O Wait)",
                "Kiểm tra hiệu năng – SQL",
                "Kiểm tra Firmware và bản vá",
                "Kiểm tra sao lưu",
                "Kiểm tra đồng bộ Active Dataguard",
                "Total"
            ]
        else:
            title = "SUMMARY"
            headers = ["NO", "CRITERIA", "EVALUATION SCORE", "EVALUATION SCORE (PREVIOUS REPORT)"]
            items = [
                "Status check",
                "Performance check – CPU",
                "Performance check – Memory",
                "Performance check – Storage (Free capacity) for each volume",
                "Performance check – Storage (I/O check) (Average/max)",
                "Performance check – SQL Execution",
                "Firmware and Patch check",
                "Backup check",
                "Standby check",
                "Total"
            ]
            
        self.doc.add_heading(title, level=3)
        
        # Static table: 11 rows (1 header + 10 data), 4 columns
        table = self.doc.add_table(rows=11, cols=4)
        table.style = 'Table Grid'
        table.allow_autofit = False
        
        # Header
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        self._format_header_row(table.rows[0])
        
        for i, item_text in enumerate(items):
            row_idx = i + 1
            cells = table.rows[row_idx].cells
            
            # STT (1-9), Empty for Total
            if i < 9:
                cells[0].text = str(row_idx)
            else:
                cells[0].text = "" # No STT for Total
                
            cells[1].text = item_text
            
            # Columns 3 & 4 empty
            cells[2].text = ""
            cells[3].text = ""
            
            # Last row styling (Total)
            if i == 9:
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                
                for cell in cells:
                    # Shading D9D9D9
                    shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._element.get_or_add_tcPr().append(shd)
                    
                    # Bold text
                    for paragraph in cell.paragraphs:
                        if not paragraph.runs:
                            run = paragraph.add_run()
                        else:
                            run = paragraph.runs[0]
                        run.bold = True
            
            # Alignment
            for j, cell in enumerate(cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Center STT, Evaluation columns
                if j != 1:
                    if cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fixed widths: 1.4, 10.0, 2.5, 2.5 (Total 16.4)
        self._set_column_widths(table, [Cm(1.4), Cm(10.0), Cm(2.5), Cm(2.5)])

    def _add_recommendation_section(self, findings: List[Dict[str, Any]]):
        title = "Khuyến nghị và giải pháp" if self.language == 'vi' else "Recommendations"
        self.doc.add_heading(title, level=3)
        
        # Table with 5 columns: STT, Khuyến nghị, Rủi ro, Mức độ, Phụ lục
        table = self.doc.add_table(rows=len(findings) + 1, cols=5)
        table.style = 'Table Grid'
        table.allow_autofit = False
        
        # Header
        if self.language == 'vi':
            headers = ["STT", "Khuyến nghị", "Rủi ro/ảnh hưởng", "Mức độ", "Tham khảo Phụ lục"]
        else:
            headers = ["NO", "Recommendations", "Risk/Effect", "Severity", "Appendix Reference"]
            
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        self._format_header_row(table.rows[0])
        
        # Fixed widths: 1.2, 6.55, 4.5, 1.9, 2.25 (Total 16.4)
        self._set_column_widths(table, [Cm(1.2), Cm(6.55), Cm(4.5), Cm(1.9), Cm(2.25)])
        
        # Data
        for i, f in enumerate(findings):
            row = table.rows[i+1].cells
            row[0].text = str(i + 1)
            
            if self.language == 'vi':
                row[1].text = f.get('rec_vi', '')
                row[2].text = f.get('risk_vi', '')
                sev_text = f.get('severity_vi', '')
            else:
                row[1].text = f.get('rec_en', '')
                row[2].text = f.get('risk_en', '')
                sev_text = f.get('severity_en', '')
            
            row[3].text = sev_text
            
            # Apply color and bold for Severity
            if row[3].paragraphs:
                p = row[3].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Determine styling
                is_severe = sev_text.upper() in ["CRITICAL", "NGHIÊM TRỌNG", "HIGH", "CAO"]
                is_critical = sev_text.upper() in ["CRITICAL", "NGHIÊM TRỌNG"]
                
                if is_severe:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0) # Red
                        if is_critical:
                            run.font.bold = True
            
            ref_text = f.get('appendix_ref', '')
            if self.language != 'vi':
                ref_text = ref_text.replace("Phụ lục", "Appendix")
            row[4].text = ref_text
            
            # Align first and last 2 cols
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for cell in row:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_appendix_ref_section(self, db_name: str):
        title = "Thông tin chi tiết của cơ sở dữ liệu" if self.language == 'vi' else \
                "Information of DB Health check"
        self.doc.add_heading(title, level=3)
        
        if self.language == 'vi':
            text = f"Tham khảo phụ lục báo cáo {db_name} đính kèm."
        else:
            text = f"Please refer to Section {db_name} in the Appendix attached along this Report."
        
        self.doc.add_paragraph(text)

    def _format_header_row(self, row):
        """Standard navy blue header formatting for consistency with Generator"""
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading_elm = parse_xml(r'<w:shd {} w:fill="1F3864"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.name = self.font_name

    def _format_table_row_bg(self, row, bg_color, bold=False):
        """Set background color for a whole row"""
        for cell in row.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            for para in cell.paragraphs:
                for run in para.runs:
                    if bold: run.font.bold = True
                    run.font.name = self.font_name

    def _set_column_widths(self, table, widths):
        """Force absolute fixed widths using internal XML manipulation (Twips/DXA)"""
        table.allow_autofit = False
        
        # 1. Set Table-level fixed layout and width
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Fixed layout: Forces Word to respect column sizes exactly
        layout = tblPr.xpath('w:tblLayout')
        if not layout:
            layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
            tblPr.append(layout)
        else:
            layout[0].set(qn('w:type'), 'fixed')
            
        # Total Table Width in DXA (Twips)
        # EMU / 635 = Twips
        total_emu = sum(w for w in widths)
        total_dxa = int(total_emu / 635)
        
        tblW = tblPr.xpath('w:tblW')
        if not tblW:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>')
            tblPr.append(tblW)
        else:
            tblW[0].set(qn('w:w'), str(total_dxa))
            tblW[0].set(qn('w:type'), 'dxa')

        # 2. Set Cell-level fixed widths for each column
        for i, width in enumerate(widths):
            width_dxa = int(width / 635)
            # Apply to column object first
            table.columns[i].width = width
            # Apply to each cell in the column via XML
            for cell in table.columns[i].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.xpath('w:tcW')
                if not tcW:
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
                    tcPr.append(tcW)
                else:
                    tcW[0].set(qn('w:w'), str(width_dxa))
                    tcW[0].set(qn('w:type'), 'dxa')

