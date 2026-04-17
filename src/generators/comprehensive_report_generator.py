"""
Comprehensive Healthcare Report Generator
Matches exact structure 1.0-1.10 with all required sections
"""
import logging
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import re

from ..config import TEMPLATE_DIR

logger = logging.getLogger(__name__)


class ComprehensiveHealthcareReportGenerator:
    """Generates the comprehensive health check report (DOCX)"""
    
    # Pre-compiled regex patterns for performance
    NUMERIC_RE = re.compile(r'^-?[\d,.]+\s*%?$')
    TIME_RE = re.compile(r'^[\d.]+\s*(us|ms|s)$', re.I)
    DATE_RE = re.compile(r'^[\d]{2,4}[-/][a-zA-Z\d]{2,3}[-/][\d]{2,4}')
    
    def __init__(self, output_path: str, font_option: str = 'times', db_role: str = 'primary'):
        self.output_path = output_path
        self.font_option = font_option.lower()
        self.db_role = db_role.lower()
        
        # Select template based on font option
        template_name = 'timenr_template.docx' if self.font_option == 'times' else 'calibri_template.docx'
        template_path = TEMPLATE_DIR / template_name
        
        try:
            if template_path.exists():
                self.doc = Document(str(template_path))
                logger.info(f"Initialized with template: {template_path}")
                # Remove initial empty paragraphs often found in templates
                while len(self.doc.paragraphs) > 0 and not self.doc.paragraphs[0].text.strip():
                    p = self.doc.paragraphs[0]._element
                    p.getparent().remove(p)
            else:
                logger.warning(f"Template {template_path} not found. Using default document.")
                self.doc = Document()
        except Exception as e:
            logger.error(f"Error loading template: {e}")
            self.doc = Document()

        self.font_name = 'Times New Roman' if self.font_option == 'times' else 'Calibri'
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure default fonts and styles with robust XML overrides"""
        try:
            # Set Normal style
            style = self.doc.styles['Normal']
            style.font.name = self.font_name
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 0)
            
            # Helper to force font in XML (essential for Headings)
            def force_font(style_obj, name):
                rPr = style_obj._element.rPr
                if rPr is None:
                    rPr = style_obj._element.make_element(qn('w:rPr'))
                    style_obj._element.append(rPr)
                
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}" w:eastAsia="{name}"/>')
                    rPr.append(rFonts)
                else:
                    rFonts.set(qn('w:ascii'), name)
                    rFonts.set(qn('w:hAnsi'), name)
                    rFonts.set(qn('w:cs'), name)
                    rFonts.set(qn('w:eastAsia'), name)
                    # Clear theme fonts which often cause Calibri fallback
                    for attr in ['asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme']:
                        if qn('w:'+attr) in rFonts.attrib:
                            del rFonts.attrib[qn('w:'+attr)]

            force_font(style, self.font_name)
            
            # Setup Headings
            headings = [
                ('Heading 1', 16),
                ('Heading 2', 13),
                ('Heading 3', 11)
            ]
            
            for name, size in headings:
                h = self.doc.styles[name]
                h.font.name = self.font_name
                h.font.size = Pt(size)
                h.font.bold = True
                h.font.color.rgb = RGBColor(0, 0, 0)
                force_font(h, self.font_name)
                # Ensure no space after/before for consistency
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
                if name == 'Heading 3':
                    h.font.italic = True
        except Exception as e:
            logger.warning(f"Could not setup styles: {e}")
            
    def _get_base_db_name(self, data: Dict[str, Any]) -> str:
        """Extract database name by removing trailing digits from instance name"""
        nodes = data.get('nodes', [])
        if not nodes:
            return data.get('db_name', 'Unknown').upper()
            
        inst = str(nodes[0].get('instance_name', '')).strip()
        if not inst or inst == 'NODE1':
            return data.get('db_name', 'Unknown').upper()
            
        db_name = inst
        while db_name and db_name[-1].isdigit():
            db_name = db_name[:-1]
        if self.db_role == 'standby':
            return f"{db_name.upper()}-STB"
        return db_name.upper()

    def _add_instance_name(self, text):
        """Add instance name with specific style"""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(text).upper())
        run.font.name = self.font_name
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.italic = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _format_table_header(self, row):
        """Format table header row to Navy Blue background, White text, Centered"""
        for cell in row.cells:
            # Vertical Align Center
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Set background color 1F3864 (Navy Blue)
            shading_elm = parse_xml(r'<w:shd {} w:fill="1F3864"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    
    def _add_node_images(self, node_dir: str, image_candidates: list):
        """
        Add images from node's generated_files or exawatcher_files folder - Fixed at 16.4 cm.
        Supports fallback names: item can be a string or a list of candidate strings.
        """
        # Search for OSWBB folder first, fallback to ExaWatcher folder
        img_folder = Path(node_dir) / 'generated_files'
        if not img_folder.exists():
            img_folder = Path(node_dir) / 'exawatcher_files'
            
        if not img_folder.exists():
            return
            
        for item in image_candidates:
            # Determine candidates for this image slot
            candidates = [item] if isinstance(item, str) else item
            
            found_path = None
            for cand_name in candidates:
                img_path = img_folder / cand_name
                if img_path.exists():
                    found_path = img_path
                    break
            
            if found_path:
                try:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(found_path), width=Cm(16.4))
                except Exception as e:
                    logger.warning(f"Error adding image from {found_path}: {e}")
            else:
                # Only log warning if we are not suppressing missing images
                if not getattr(self, '_hide_missing_images', False):
                    logger.warning(f"None of the candidates {candidates} were found in {img_folder}")
    
    def generate_from_parsed_data(self, data: Dict[str, Any]) -> bool:
        """Generate comprehensive report"""
        try:
            # 1.0 Database Name & Basic Info
            self._add_section_database_title(data)
            
            # 1.1 Status Check
            self._add_section_1_1_status_check(data)
            
            # 1.2 Alert Logs
            self._add_section_1_2_alert_logs(data)
            
            # 1.3 Performance Check
            self._add_section_1_3_performance_check(data)
            
            # 1.4 Data Quality (Skip if Standby)
            if self.db_role == 'primary':
                self._add_section_1_4_data_quality(data)
            
            # 1.5 HA/Clusterware Status
            self._add_section_1_5_ha_status(data)
            
            # 1.6 Storage Capacity
            self._add_section_1_6_storage_capacity(data)
            
            # 1.7 Backup Status
            self._add_section_1_7_backup_status(data)
            
            # 1.8 Dataguard Status
            self._add_section_1_8_dataguard_status(data)
            
            # 1.9 Security (Skip if Standby)
            if self.db_role == 'primary':
                self._add_section_1_9_security(data)
            
            # 1.10 Patch Update
            self._add_section_1_10_patch_update(data)
            
            self.doc.save(self.output_path)
            logger.info(f"Comprehensive report saved to {self.output_path}")
            return True
        
        except Exception as e:
            logger.error(f"Failed to generate report: {e}")
            return False
    
    def _add_section_database_title(self, data: Dict[str, Any]):
        """Database name as main title"""
        db_name = self._get_base_db_name(data)
        self.doc.add_heading(db_name, level=1)
        self.doc.add_paragraph()
    
    def _add_section_1_1_status_check(self, data: Dict[str, Any]):
        """1.1 Status Check"""
        self.doc.add_heading("Status check", level=2)
        

        db_name = self._get_base_db_name(data)
        headers = ["DB NAME", "LISTENER STATUS", "DB STATUS", "GRID CLUSTER", "ASM"]
        values = [db_name, "Running", "Open", "Online", "Running"]
        self._create_table_from_rows([headers, values])
    
    def _add_section_1_2_alert_logs(self, data: Dict[str, Any]):
        """1.2 Alert Logs"""
        self.doc.add_heading("Alert log", level=2)
        
        nodes = data.get('nodes', [])
        for node in nodes:
            instance_name = node.get('instance_name', f"NODE{node.get('node_id')}")
            self._add_instance_name(instance_name)
            
            alerts = list(node.get('alerts', {}).get('alerts', []))
            alerts.sort(key=lambda x: str(x.get('timestamp', '')), reverse=True)
            self._add_alert_table(alerts)
    
    def _add_alert_table(self, alerts: List[Dict[str, Any]]):
        """Add alert table for a node"""
        if not alerts:
            p = self.doc.add_paragraph("No alerts in the specified period")
            return
        
        # Data rows
        for alert in alerts:
            row_data = [
                str(alert.get('timestamp', '')).replace('T', ' ')[:19],
                str(alert.get('full_text', alert.get('error_code', '')))
            ]
            self._create_table_from_rows([hdr_text, row_data], col_widths=[Cm(4.0), Cm(12.4)], max_rows=None)
            # Remove the header from subsequent rows in a better refactor
        # Let's actually refactor _add_alert_table to be more standard
        pass

    def _add_alert_table(self, alerts: List[Dict[str, Any]]):
        """Add alert table for a node - Refactored to use 16.4cm"""
        if not alerts:
            self.doc.add_paragraph("No alerts in the specified period")
            return
        
        rows = [["DATE TIME", "ERROR"]]
        for alert in alerts:
            rows.append([
                str(alert.get('timestamp', '')).replace('T', ' ')[:19],
                str(alert.get('full_text', alert.get('error_code', '')))
            ])
            
        self._create_table_from_rows(rows, max_rows=30, col_widths=[Cm(4.0), Cm(12.4)], 
                                    truncate_length=None, align_left_cols=[0])
    
    def _add_section_1_3_performance_check(self, data: Dict[str, Any]):
        """1.3 Performance Check with subsections"""
        self.doc.add_heading("Performance check", level=2)
        
        nodes = data.get('nodes', [])
        
        # 1.3.1 CPU
        self.doc.add_heading("CPU", level=3)
        for node in nodes:
            inst = node.get('instance_name', f"NODE{node.get('node_id')}")
            self._add_instance_name(inst)
            cpu_candidates = [
                ['OSWg_OS_Cpu_Idle.jpg', 'OSWg_Cpu_Idle.jpg', 'Exa_Cpu_Idle.png'],
                ['OSWg_OS_Cpu_System.jpg', 'OSWg_Cpu_System.jpg', 'Exa_Cpu_Sys.png'],
                ['OSWg_OS_Cpu_User.jpg', 'OSWg_Cpu_User.jpg', 'Exa_Cpu_Usr.png']
            ]
            self._add_node_images(node.get('data_dir', ''), cpu_candidates)
        
        # 1.3.2 Memory
        self.doc.add_heading("Memory", level=3)
        for node in nodes:
            inst = node.get('instance_name', f"NODE{node.get('node_id')}")
            self._add_instance_name(inst)
            mem_candidates = [
                ['OSWg_OS_Memory_Free.jpg', 'OSWg_MemInfoFreeMem.jpg', 'Exa_Mem_OS.png'],
                ['OSWg_OS_Memory_Swap.jpg', 'OSWg_Memory_Swap.jpg', 'Exa_Mem_HugePages.png']
            ]
            self._add_node_images(node.get('data_dir', ''), mem_candidates)
            self.doc.add_paragraph("- Buffer & Library Hit Ratio")
            self._add_efficiency_table(node.get('awr', {}), align_left_cols=[0, 2])
        
        # 1.3.3 I/O
        self.doc.add_heading("I/O", level=3)
        for node in nodes:
            inst = node.get('instance_name', f"NODE{node.get('node_id')}")
            self._add_instance_name(inst)
            io_candidates = [
                ['OSWg_OS_IO_PB.jpg', 'OSWg_IO_PB.jpg', 'Exa_IO_Summary.png']
            ]
            self._add_node_images(node.get('data_dir', ''), io_candidates)
            
            if self.db_role == 'primary':
                self.doc.add_paragraph("- Wait Classes by Total Wait Time")
                self._add_awr_table(node.get('awr', {}), 'wait class', align_left_cols=[0])
        
        if self.db_role == 'primary':
            # 1.3.4 Top Queries
            self.doc.add_heading("Top queries", level=3)
            # 1.3.4: Fixed CM widths per user request
            top_sql_widths = [Cm(1.9), Cm(1.5), Cm(2.25), Cm(1.25), Cm(1.25), Cm(1.25), Cm(3.0), Cm(4.0)]
            for node in nodes:
                inst = node.get('instance_name', f"NODE{node.get('node_id')}")
                self._add_instance_name(inst)
                self._add_awr_table(node.get('awr', {}), 'top SQL elapsed', drop_cols=['SQL Module'], col_widths=top_sql_widths)
            
            # 1.3.5 SQL Text
            self.doc.add_heading("Lists SQL Text", level=3)
            for node in nodes:
                inst = node.get('instance_name', f"NODE{node.get('node_id')}")
                self._add_instance_name(inst)
                self._add_filtered_sql_text_table(node.get('awr', {}))
            
            # 1.3.6 Wait Events
            self.doc.add_heading("Top wait events", level=3)
            # 1.3.6: Fixed CM widths per user request
            wait_event_widths = [Cm(5.2), Cm(2.0), Cm(2.5), Cm(2.0), Cm(1.7), Cm(3.0)]
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
            for node in nodes:
                inst = node.get('instance_name', f"NODE{node.get('node_id')}")
                self._add_instance_name(inst)
                table = self._add_awr_table(node.get('awr', {}), 'Top 10 Foreground Events by Total Wait Time', col_widths=wait_event_widths)
                
                # Formating column 6 logic (Align Center Left)
                if table and len(table.columns) >= 6:
                    for row in table.rows[1:]:
                        if len(row.cells) >= 6:
                            cell = row.cells[5]
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        db_info_first = nodes[0].get('database_info', {}) if nodes else {}
        
        # 1.3.7 Disk group usage - Col 2 center
        self.doc.add_heading("Disk group usage", level=3)
        self._add_db_info_table(db_info_first, 'ASM', align_center_cols=[1], filter_nulls=True)
        
        # 1.3.8 Tablespace usage - Fixed CM widths
        self.doc.add_heading("Tablespace usage", level=3)
        tablespace_widths = [Cm(4.15), Cm(1.75), Cm(2.25), Cm(2.25), Cm(2.25), Cm(2.25), Cm(1.5)]
        self._add_db_info_table(db_info_first, 'TABLESPACE', col_widths=tablespace_widths)
        
        if self.db_role == 'primary':
            self.doc.add_heading("Index Fragment", level=3)
            self.doc.add_paragraph("- List of normal index fragment")
            frag_widths_norm = [Cm(2.4), Cm(8.0), Cm(2.0), Cm(2.0), Cm(2.0)]
            self._add_db_info_table(db_info_first, 'INDEX_FRAGMENT', filter_nulls=True, col_widths=frag_widths_norm)
            self.doc.add_paragraph("- List of partition index fragment")
            frag_widths_part = [Cm(2.4), Cm(4.0), Cm(4.0), Cm(2.0), Cm(2.0), Cm(2.0)]
            self._add_db_info_table(db_info_first, 'INDEX_PARTITION_FRAGMENT', filter_nulls=True, col_widths=frag_widths_part)
            
            # 1.3.10 Table Fragment
            self.doc.add_heading("Table Fragment", level=3)
            self.doc.add_paragraph("- List of normal table fragment")
            self._add_db_info_table(db_info_first, 'TABLE_FRAGMENT', filter_nulls=True, col_widths=frag_widths_norm)
            self.doc.add_paragraph("- List of partition table fragment")
            self._add_db_info_table(db_info_first, 'TABLE_PARTITION_FRAGMENT', filter_nulls=True, col_widths=frag_widths_part)
    
    def _add_section_1_4_data_quality(self, data: Dict[str, Any]):
        """1.4 Data Quality"""
        self.doc.add_heading("Data quality", level=2)
        
        nodes = data.get('nodes', [])
        db_info_first = nodes[0].get('database_info', {}) if nodes else {}
        
        # 1.4.1 Invalid Objects
        self.doc.add_heading("Invalid Object", level=3)
        self._add_db_info_table(db_info_first, 'INVALID_OBJECT', filter_nulls=True, 
                               align_center_cols=[1], align_left_cols=[2])
        
        # 1.4.2 Table/Index Statistics
        self.doc.add_heading("Tables/Indexes Statistics", level=3)
        
        # Table Statistics
        table_stats = db_info_first.get('TABLE_STATISTICS', [])
        if table_stats and len(table_stats) > 1:
            header = [str(c).upper() for c in table_stats[0]]
            date_col = -1
            for i, c in enumerate(header):
                if 'ANALYZED' in c or 'DATE' in c:
                    date_col = i
                    break
            last_analyzed = str(table_stats[1][date_col]).strip().split(" ")[0] if date_col != -1 and len(table_stats[1]) > date_col else "Unknown"
            self.doc.add_paragraph(f"The tables was last analyzed on {last_analyzed}")
        self._add_db_info_table(db_info_first, 'TABLE_STATISTICS', col_widths=[Cm(3.7), Cm(10.0), Cm(2.7)])
        
        # Index Statistics
        index_stats = db_info_first.get('INDEX_STATISTICS', [])
        if index_stats and len(index_stats) > 1:
            header = [str(c).upper() for c in index_stats[0]]
            date_col = -1
            for i, c in enumerate(header):
                if 'ANALYZED' in c or 'DATE' in c:
                    date_col = i
                    break
            last_analyzed = str(index_stats[1][date_col]).strip().split(" ")[0] if date_col != -1 and len(index_stats[1]) > date_col else "Unknown"
            self.doc.add_paragraph(f"The indexes was last analyzed on {last_analyzed}")
        self._add_db_info_table(db_info_first, 'INDEX_STATISTICS', col_widths=[Cm(3.7), Cm(5.0), Cm(5.0), Cm(2.7)])
    
    def _add_section_1_5_ha_status(self, data: Dict[str, Any]):
        """1.5 HA/Clusterware Status"""
        self.doc.add_heading("HA/Clusterware status", level=2)
        
        nodes = data.get('nodes', [])
        
        # 1.5.1 Clusterware Status
        self.doc.add_heading("Clusterware status", level=3)
        combined_cluster = []
        for node in nodes:
            cluster_node = node.get('database_info', {}).get('CHECK_CLUSTER', [])
            if cluster_node:
                start_idx = 1 if len(combined_cluster) > 0 else 0
                combined_cluster.extend(cluster_node[start_idx:])
                
        if combined_cluster:
            # 1.5.1: Fixed CM widths
            cluster_widths = [Cm(6.4), Cm(10.0)]
            self._create_table_from_rows(combined_cluster, max_rows=40, col_widths=cluster_widths, truncate_length=None)
        else:
            self.doc.add_paragraph("[Section 'CHECK_CLUSTER' not found]")
            
        # 1.5.2 CRS Resources - Only Node 1 per user request
        self.doc.add_heading("Clusterware services status details", level=3)
        combined_res = []
        if nodes:
            res_n = nodes[0].get('database_info', {}).get('RESOURCE_CRS', [])
            if res_n:
                combined_res.extend(res_n)
                
        if combined_res:
             # 1.5.2: Fixed CM widths
             res_widths = [Cm(5.65), Cm(2.0), Cm(2.0), Cm(3.0), Cm(3.75)]
             self._create_table_from_rows(combined_res, max_rows=100, col_widths=res_widths)
        else:
             self.doc.add_paragraph("[Section 'RESOURCE_CRS' not found]")
    
    def _add_section_1_6_storage_capacity(self, data: Dict[str, Any]):
        """1.6 Storage Capacity"""
        self.doc.add_heading("Storage capacity", level=2)
        
        nodes = data.get('nodes', [])
        for node in nodes:
            inst = node.get('instance_name', f"NODE{node.get('node_id')}")
            self._add_instance_name(inst)
            
            table_info = node.get('database_info', {}).get('DISK_USAGE')
            if not table_info:
                # Fallback search
                for other_node in nodes:
                    table_info = other_node.get('database_info', {}).get('DISK_USAGE')
                    if table_info: break
                    
            proxy_db = {}
            if table_info:
                proxy_db['DISK_USAGE'] = table_info
            
            table_out = self._add_db_info_table(proxy_db, 'DISK_USAGE')
            self._scale_storage_columns(table_out)
        
    def _scale_storage_columns(self, table):
        """Scale columns for 1.6 Storage table - Enforced 16.4 cm"""
        if table and len(table.columns) >= 6:
            # Total 16.4cm: 6.0 (Label), 1.5, 1.5, 1.5, 1.5, 4.4 (Mount)
            widths = [Cm(6.0), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(4.4)]
            self._set_cell_widths(table, widths)
            
            # Align Col 2, 3, 4 to Right (indices 1, 2, 3)
            for row in table.rows[1:]: # Skip header
                for idx in [1, 2, 3]:
                    if idx < len(row.cells):
                        for p in row.cells[idx].paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Removed redundant add_paragraph() call source-level spacing is handled by _create_table_from_rows or Heading space_before
    
    def _add_section_1_7_backup_status(self, data: Dict[str, Any]):
        """1.7 Backup Status"""
        self.doc.add_heading("Backup status", level=2)
        nodes = data.get('nodes', [])
        db_info_first = nodes[0].get('database_info', {}) if nodes else {}
        
        # 1.7.1 Backup Status
        self.doc.add_heading("Backup status", level=3)
        
        table_data = db_info_first.get('CHECK_BACKUP', [])
        if table_data and len(table_data) > 1:
            header = [str(c).upper() for c in table_data[0]]
            status_idx = -1
            for i, c in enumerate(header):
                if 'STATUS' in c:
                    status_idx = i
                    break
            
            filtered = [table_data[0]]
            for row in table_data[1:]:
                # Ignore trailing empty rows early
                if all(not str(cell).strip() for cell in row):
                    continue
                if status_idx != -1 and status_idx < len(row):
                    if str(row[status_idx]).strip().upper() == 'NULL':
                        continue
                filtered.append(row)
                
            if filtered and len(filtered) > 0:
                # Rename headers for Col 5 and 6 (Index 4 and 5)
                if len(filtered[0]) >= 6:
                    filtered[0][4] = "IN GBs"
                    filtered[0][5] = "OUT GBs"
                
            if len(filtered) == 1:
                filtered.append([""] * len(filtered[0]))
            
            backup_widths = [Cm(2.0), Cm(2.0), Cm(1.75), Cm(2.25), Cm(1.25), Cm(1.25), Cm(3.0), Cm(2.9)]
            self._create_table_from_rows(filtered, max_rows=None, col_widths=backup_widths, font_size=9)
        else:
            backup_widths = [Cm(2.0), Cm(2.0), Cm(1.75), Cm(2.25), Cm(1.25), Cm(1.25), Cm(3.0), Cm(2.9)]
            self._add_db_info_table(db_info_first, 'CHECK_BACKUP', col_widths=backup_widths, font_size=9)
        
        # 1.7.2 Scheduling
        self.doc.add_heading("Scheduling", level=3)
        p = self.doc.add_paragraph("• Backup level 0:")
        p.paragraph_format.left_indent = Inches(0.25)
        p = self.doc.add_paragraph("Monday to Sunday (about time: 19:50 ~ 20:55)")
        p.paragraph_format.left_indent = Inches(0.5)
        
        p = self.doc.add_paragraph("• Backup archive:")
        p.paragraph_format.left_indent = Inches(0.25)
        p = self.doc.add_paragraph("Monday to Sunday (about time: 06:00; 14:00)")
        p.paragraph_format.left_indent = Inches(0.5)
        
        # 1.7.3 Policy
        self.doc.add_heading("Policy", level=3)
        p = self.doc.add_paragraph("Backups work with the backup policy:")
        
        backup_policy = db_info_first.get('BACKUP_POLICY', [])
        if backup_policy and len(backup_policy) > 1:
            row_idx = 1
            if row_idx < len(backup_policy):
                policy_text = ""
                # Find the row starting with 'configure' in 2nd row (index 1)
                for cell in backup_policy[row_idx]:
                    if 'configure' in str(cell).lower():
                        policy_text = str(cell)
                        break
                if not policy_text and len(backup_policy[row_idx]) > 0:
                     policy_text = backup_policy[row_idx][0]
                if policy_text:
                    self.doc.add_paragraph(policy_text)
    
    def _add_section_1_8_dataguard_status(self, data: Dict[str, Any]):
        """1.8 Dataguard Status"""
        self.doc.add_heading("Dataguard status", level=2)
        
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        headers = ["CONFIGURATION STATUS", "APPLY MODE", "APPLY LAG (SECONDS)", "TRANSPORT LAG (SECONDS)"]
        values = ["SUCCESS", "REAL TIME", "0", "0"]
        
        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        self._format_table_header(table.rows[0])
        
        # Align row 2 (data) and setup fonts
        for i, value in enumerate(values):
            cell = table.rows[1].cells[i]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = self.font_name
                    run.font.size = Pt(10)
        
        # Enforce 16.4cm width
        self._set_cell_widths(table, [Cm(4.1), Cm(4.1), Cm(4.1), Cm(4.1)])
    
    def _add_section_1_9_security(self, data: Dict[str, Any]):
        """1.9 Security"""
        self.doc.add_heading("Security", level=2)
        
        nodes = data.get('nodes', [])
        db_info_first = nodes[0].get('database_info', {}) if nodes else {}
        
        # 1.9.1 DBA Users
        self.doc.add_heading("DBA users", level=3)
        p = self.doc.add_paragraph("List user who has DBA privilege others than SYS/SYSTEM")
        self._add_db_info_table(db_info_first, 'DBA_ROLE', filter_nulls=True)
        
        # 1.9.2 Objects in SYSTEM
        self.doc.add_heading("Users With Objects in Tablespace SYSTEM/SYSAUX", level=3)
        
        obj_sys = db_info_first.get('OBJECT_IN_SYSTEM', [])
        if obj_sys:
            filtered_sys = [obj_sys[0]]
            for row in obj_sys[1:]:
                row_str = " ".join([str(c).strip() for c in row]).upper()
                if any(p in row_str for p in ["NULL NULL", "NULL 0", "NULL | 0", "NULL  0"]):
                    continue
                filtered_sys.append(row)
            if len(filtered_sys) == 1:
                # Add a blank row to keep the frame
                filtered_sys.append([""] * len(obj_sys[0]))
            self._create_table_from_rows(filtered_sys, max_rows=20, col_widths=[Cm(3.5), Cm(7.0), Cm(3.25), Cm(2.65)])
        else:
            # Create even if section missing
            self.doc.add_paragraph("[Section 'OBJECT_IN_SYSTEM' not found or empty]")
            headers = ["OWNER", "OBJECT TYPE", "TABLESPACE", "COUNT"]
            self._create_table_from_rows([headers, ["", "", "", ""]], col_widths=[Cm(3.5), Cm(7.0), Cm(3.25), Cm(2.65)])
    
    def _add_section_1_10_patch_update(self, data: Dict[str, Any]):
        """1.10 Patch Update Status"""
        self.doc.add_heading("Patch update status", level=2)
        
        nodes = data.get('nodes', [])
        combined_patch = []
        for node in nodes:
            patch_node = node.get('database_info', {}).get('CHECK_PATCHES', [])
            if patch_node:
                start_idx = 1 if len(combined_patch) > 0 else 0
                combined_patch.extend(patch_node[start_idx:])
                
        if combined_patch and len(combined_patch) > 1:
            for i in range(1, len(combined_patch)):
                if len(combined_patch[i]) > 0:
                    server_name = str(combined_patch[i][0])
                    if "." in server_name:
                        combined_patch[i][0] = server_name.split(".")[0]
                        
            table = self._create_table_from_rows(combined_patch, max_rows=None, truncate_length=None,
                                              col_widths=[Cm(3.2), Cm(3.2), Cm(10.0)])
        else:
            self.doc.add_paragraph("[Section 'CHECK_PATCHES' not found]")
    
    def _add_awr_table(self, awr_node_data: Dict, table_name: str, drop_cols: List[str] = None,
                      align_left_cols: List[int] = None, col_widths: List[float] = None):
        """Add table from AWR data - search by keywords"""
        node_tables = awr_node_data.get('tables', [])
        
        # Create search keywords
        search_terms = []
        
        # Add original name first (highest priority)
        term_clean = table_name.lower().strip()
        search_terms.append(term_clean)
        
        # If it's a multi-word name, add it without common filler words
        keywords = term_clean.split()
        if len(keywords) > 2:
             # Try combining key technical terms
             technical_terms = [k for k in keywords if k not in ['by', 'total', 'the', 'of', 'in', 'and']]
             if len(technical_terms) > 1:
                 search_terms.append(' '.join(technical_terms))
        
        # Finally add single keywords as fallback, but avoid too common ones
        for k in keywords:
            if k not in ['wait', 'events', 'total', 'time', 'by', 'the', 'of'] and len(k) > 3:
                search_terms.append(k)
        
        # Special case for 1.3.6 to ensure it hits the right Foreground Events table
        # Matches common AWR summary attributes for this specific section
        if 'foreground events' in term_clean or 'wait events' in term_clean:
            search_terms = [
                'top 10 foreground events', 
                'top 5 timed foreground events',
                'this table displays top 10 wait events by total wait time',
                'foreground events'
            ]

        for table_info in node_tables:
            title = table_info.get('title', '').lower()
            # Check if any search term is in the table title
            for search_term in search_terms:
                if search_term.strip() and search_term.strip() in title:
                    rows = table_info.get('rows', [])
                    if rows:
                        if drop_cols:
                            header = [str(c).upper() for c in rows[0]]
                            indices_to_drop = []
                            for dc in drop_cols:
                                for idx, h in enumerate(header):
                                    if dc.upper() in h:
                                        indices_to_drop.append(idx)
                            
                            if indices_to_drop:
                                filtered_rows = []
                                for r in rows:
                                    filtered_row = [v for idx, v in enumerate(r) if idx not in indices_to_drop]
                                    filtered_rows.append(filtered_row)
                                rows = filtered_rows
                        
                        return self._create_table_from_rows(rows, max_rows=15, align_left_cols=align_left_cols, col_widths=col_widths)
                    return None
        
        # Not found - log what we were looking for
        logger.warning(f"Table '{table_name}' not found in AWR data for node")
        if not getattr(self, '_hide_missing_awr', False):
            self.doc.add_paragraph(f"[Table '{table_name}' not found in AWR data]")
        return None
    
    def _add_filtered_sql_text_table(self, awr_node_data: Dict):
        """Extract SQL IDs from Top SQL and filter Complete List of SQL Text"""
        node_tables = awr_node_data.get('tables', [])
        
        # 1. Find Top SQL table
        top_sql_rows = []
        for table_info in node_tables:
            title = table_info.get('title', '').lower()
            if 'top sql by elapsed time' in title:
                top_sql_rows = table_info.get('rows', [])
                break
                
        # Extract SQL IDs from Top SQL (column 6 is SQL Id)
        sql_ids = set()
        sql_ids_ordered = []
        if top_sql_rows and len(top_sql_rows) > 0:
            # Skip header row(s)
            for row in top_sql_rows[1:]:
                if len(row) > 6:
                    sql_id = str(row[6]).strip()
                    if sql_id and sql_id not in sql_ids:
                        sql_ids.add(sql_id)
                        sql_ids_ordered.append(sql_id)
        
        # 2. Find Complete SQL Text table
        sql_text_rows = []
        for table_info in node_tables:
            title = table_info.get('title', '').lower()
            if 'text of the sql statements' in title or 'complete list of sql text' in title:
                sql_text_rows = table_info.get('rows', [])
                break
                
        if not sql_text_rows:
            self.doc.add_paragraph("[Table 'Complete List of SQL Text' not found in AWR data]")
            return
            
        # Map out SQL texts
        sql_texts_map = {}
        if sql_text_rows:
            for row in sql_text_rows[1:]:
                if len(row) > 0:
                    current_id = str(row[0]).strip()
                    sql_texts_map[current_id] = row
                    
        filtered_rows = [sql_text_rows[0]] if sql_text_rows else [["SQL ID", "SQL Text"]]
        for sid in sql_ids_ordered:
            if sid in sql_texts_map:
                filtered_rows.append(sql_texts_map[sid])
                        
        if len(filtered_rows) > 1:
            self._create_table_from_rows(filtered_rows, max_rows=None, truncate_length=None, 
                                        col_widths=[Cm(4.0), Cm(12.4)])
        else:
            # If no SQL text, add a blank frame
            self.doc.add_paragraph("No matching SQL texts found for Top SQLs.")
            self._create_table_from_rows([filtered_rows[0], ["", ""]], col_widths=[Cm(4.0), Cm(12.4)])
    
    def _add_db_info_table(self, db_info: Dict, section_key: str, filter_nulls: bool = False,
                          align_left_cols: List[int] = None, align_center_cols: List[int] = None, 
                          col_widths: List[float] = None, font_size: int = 10):
        """Add table from database_information.html"""
        table_data = db_info.get(section_key, [])
        
        if not table_data:
            self.doc.add_paragraph(f"[Section '{section_key}' not found]")
            return None
        
        # Remove empty rows or NULL rows at the end
        while table_data and len(table_data) > 1:
            row_str = " ".join([str(c).strip() for c in table_data[-1]]).upper()
            null_patterns = ["NULL NULL", "NULL 0 NULL", "NULL NULL NULL", "NULL NULL NULL 0", "NULL 0", "NULL | 0", "NULL  0"]
            is_null_row = any(p in row_str or row_str == p for p in null_patterns)
            
            if is_null_row or all(not str(cell).strip() for cell in table_data[-1]):
                table_data.pop()
            else:
                break

        if filter_nulls:
            if len(table_data) > 1:
                filtered = [table_data[0]] # Keep header
                for row in table_data[1:]:
                    row_str = " ".join([str(c).strip() for c in row]).upper()
                    # Patterns to skip
                    null_patterns = ["NULL NULL", "NULL 0 NULL", "NULL NULL NULL", "NULL NULL NULL 0", "NULL 0", "NULL | 0", "NULL  0"]
                    should_skip = False
                    for pattern in null_patterns:
                        if pattern in row_str or row_str == pattern:
                            should_skip = True
                            break
                    if not should_skip:
                        filtered.append(row)
                table_data = filtered
        
        # If table only has header row initially (or after filtering), add empty row
        if len(table_data) == 1:
            table_data.append([""] * len(table_data[0]))
        
        return self._create_table_from_rows(table_data, max_rows=None, align_left_cols=align_left_cols, 
                                          align_center_cols=align_center_cols, col_widths=col_widths,
                                          font_size=font_size)
    
    def _add_efficiency_table(self, awr_node_data: Dict, align_left_cols: List[int] = None):
        """Format 1.3.2 Instance Efficiency Table specifically"""
        node_tables = awr_node_data.get('tables', [])
        target_table = None
        for t in node_tables:
            if 'instance efficiency' in t.get('title', '').lower():
                target_table = t
                break
                
        if not target_table:
            self.doc.add_paragraph("[Table 'instance efficiency' not found in AWR data]")
            return
            
        rows = target_table.get('rows', [])
        custom_header = ["INSTANCE EFFICIENCY PERCENTAGES (TARGET 100%)", "", "", ""]
        filtered_rows = [custom_header]
        for r in rows:
            while len(r) < 4:
                r.append("")
            filtered_rows.append(r[:4])
            
        table = self._create_table_from_rows(filtered_rows, align_left_cols=align_left_cols)
        # Span custom header
        table.rows[0].cells[0].merge(table.rows[0].cells[3])
        table.rows[0].cells[0].text = custom_header[0]
        
        # Specific widths: Target 16.4cm -> 5.9cm (label), 2.3cm (value) x 2
        eff_widths = [Cm(5.9), Cm(2.3), Cm(5.9), Cm(2.3)]
        
        # Format header string (Center Left alignment)
        for p in table.rows[0].cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Force widths using cell-level logic
        self._set_cell_widths(table, eff_widths)
        
        for i, row in enumerate(table.rows):
            if i > 0:
                for p in row.cells[1].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for p in row.cells[3].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _set_cell_widths(self, tbl, widths):
        """Helper to force widths efficiently - Sets grid and first row only"""
        # 1. Set Table Grid (W:tblGrid) - Essential for layout
        tbl_pr = tbl._element.xpath('w:tblPr')[0]
        
        # Ensure w:tblLayout is fixed
        existing_layout = tbl_pr.xpath('w:tblLayout')
        if not existing_layout:
            tbl_layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
            tbl_pr.append(tbl_layout)
        else:
            existing_layout[0].set(qn('w:type'), 'fixed')

        # Set Table Width
        existing_w = tbl_pr.xpath('w:tblW')
        total_twips = sum(w.twips for w in widths)
        if not existing_w:
            tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_twips}" w:type="dxa"/>')
            tbl_pr.append(tbl_w)
        else:
            existing_w[0].set(qn('w:w'), str(total_twips))
            existing_w[0].set(qn('w:type'), 'dxa')

        # 2. Set Grid Columns (Direct XML for global structure)
        grid_xml = f'<w:tblGrid {nsdecls("w")}>'
        for w in widths:
            grid_xml += f'<w:gridCol w:w="{w.twips}"/>'
        grid_xml += '</w:tblGrid>'
        
        # Insert tblGrid after tblPr (index 0)
        tbl_grid = tbl._element.xpath('w:tblGrid')
        if tbl_grid:
            tbl._element.remove(tbl_grid[0])
        tbl._element.insert(1, parse_xml(grid_xml))

        # 3. Set widths for EVERY cell (Restoring "Hard Format" but efficiently)
        # We use table._cells flat list to make it fast O(N*M) but with low overhead
        all_cells = tbl._cells
        num_cols = len(widths)
        for idx, cell in enumerate(all_cells):
            col_idx = idx % num_cols
            if col_idx < num_cols:
                cell.width = widths[col_idx]
                        
    def _create_table_from_rows(self, rows: List[List[str]], max_rows: int = 20, truncate_length: int = 100,
                               align_left_cols: List[int] = None, align_center_cols: List[int] = None,
                               align_right_cols_forced: List[int] = None, col_widths: List[float] = None,
                               font_size: int = 10):
        """Create DOCX table from row data - Optimized for large datasets"""
        if not rows:
            return None
        
        if max_rows is not None:
            rows = rows[:max_rows]
            
        num_rows = len(rows)
        num_cols = max(len(row) for row in rows) if rows else 1
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Force column widths efficiently
        if col_widths and len(col_widths) >= num_cols:
            calc_widths = col_widths[:num_cols]
        else:
            calc_widths = [Cm(16.4 / num_cols)] * num_cols
            
        self._set_cell_widths(table, calc_widths)

        # Pre-determine alignment strategies per column based on header
        col_align_strategies = [] # List of (align_right, align_center)
        header = [str(c).upper() for c in rows[0]]
        
        for j in range(num_cols):
            h_text = header[j] if j < len(header) else ""
            align_right = any(term in h_text for term in ["WAIT", "TIME", "COUNT", "SIZE", "PERCENT", "%", "DATE", "VAL", "WAITS"])
            align_center = False
            
            if align_left_cols and j in align_left_cols:
                align_right = False
            if align_center_cols and j in align_center_cols:
                align_right = False
                align_center = True
            if align_right_cols_forced and j in align_right_cols_forced:
                align_right = True
                align_center = False
            
            col_align_strategies.append((align_right, align_center))

        # Fill table using flat cells list for much faster access
        all_cells = table._cells
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j >= num_cols: continue
                
                idx = i * num_cols + j
                cell = all_cells[idx]
                
                text_val = str(cell_text)
                if truncate_length is not None:
                    text_val = text_val[:truncate_length]
                
                cell.text = text_val
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Apply alignment
                para = cell.paragraphs[0]
                if i == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    align_right, align_center = col_align_strategies[j]
                    
                    # FINE-GRAINED NUMERIC CHECK: Only if not explicitly aligned by caller
                    is_explicitly_aligned = (align_left_cols and j in align_left_cols) or \
                                            (align_center_cols and j in align_center_cols) or \
                                            (align_right_cols_forced and j in align_right_cols_forced)

                    if not is_explicitly_aligned:
                        val_clean = text_val.strip()
                        if len(val_clean) <= 20 or " " not in val_clean:
                            if self.NUMERIC_RE.match(val_clean) or self.TIME_RE.match(val_clean) or self.DATE_RE.match(val_clean):
                                align_right = True
                    
                    if align_center:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align_right:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Set font
                for run in para.runs:
                    run.font.name = self.font_name
                    run.font.size = Pt(font_size)
        
        # Restore header formatting (Navy Blue, White text, Bold)
        if num_rows > 0:
            self._format_table_header(table.rows[0])
            
        # Add spacing after table
        self.doc.add_paragraph()
        
        return table
    
    def save(self) -> bool:
        """Save document"""
        try:
            self.doc.save(self.output_path)
            return True
        except Exception as e:
            logger.error(f"Failed to save: {e}")
            return False
